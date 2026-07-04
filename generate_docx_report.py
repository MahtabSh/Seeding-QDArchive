#!/usr/bin/env python3
"""
Generate Word (.docx) classification reports per repository for SQ26.

Each docx mirrors the 3-page PDF report:
  - Section 1: ISIC section histogram (chart embedded as image)
  - Section 2: Top-20 class table + written comments
  - Section 3: Division horizontal bar chart (embedded as image)

Output: 23692652-sq26-report-repo{id}.docx  (one per repository)

Usage:
    python3 generate_docx_report.py
    python3 generate_docx_report.py --repo 1 10 19
"""
from __future__ import annotations

import argparse
import io
import sqlite3
import textwrap
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.gridspec as gridspec
from matplotlib.patches import Patch
import numpy as np

try:
    from docx import Document
    from docx.shared import Pt, Cm, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Install: pip install python-docx")
    raise

DB = "23692652-sq26.db"

ISIC_SECTIONS: dict[str, str] = {
    "A": "Agriculture, forestry and fishing",
    "B": "Mining and quarrying",
    "C": "Manufacturing",
    "D": "Electricity, gas, steam and air conditioning supply",
    "E": "Water supply; sewerage, waste management",
    "F": "Construction",
    "G": "Wholesale and retail trade",
    "H": "Transportation and storage",
    "I": "Accommodation and food service activities",
    "J": "Publishing, broadcasting, and content production",
    "K": "ICT (telecom, programming, computing infrastructure)",
    "L": "Financial and insurance activities",
    "M": "Real estate activities",
    "N": "Professional, scientific and technical activities",
    "O": "Administrative and support service activities",
    "P": "Public administration and defence",
    "Q": "Education",
    "R": "Human health and social work activities",
    "S": "Arts, sports and recreation",
    "T": "Other service activities",
    "U": "Households – goods/services for own use",
    "V": "Extraterritorial organizations and bodies",
}

SHORT_LABELS: dict[str, str] = {
    "A": "Agri.", "B": "Mining", "C": "Mfg.", "D": "Energy",
    "E": "Water", "F": "Constr.", "G": "Trade", "H": "Transport",
    "I": "Food/Hotel", "J": "Publish.", "K": "ICT", "L": "Finance",
    "M": "Real Est.", "N": "Prof./Sci.", "O": "Admin.", "P": "Gov.",
    "Q": "Educ.", "R": "Health", "S": "Arts", "T": "Other Svc.",
    "U": "Household", "V": "Intl. Org.",
}

REPO_NAMES: dict[int, str] = {
    1:  "Zenodo",
    10: "Harvard Dataverse",
    19: "Columbia University Libraries",
}

BLUE_DARK  = "#1565C0"
BLUE_LIGHT = "#90CAF9"
GRID_COLOR = "#E0E0E0"

_SEC_COLORS: dict[str, str] = {
    "A": "#2E7D32", "B": "#4E342E", "C": "#E65100", "D": "#F9A825",
    "E": "#00838F", "F": "#546E7A", "G": "#AD1457", "H": "#6A1B9A",
    "I": "#880E4F", "J": "#1565C0", "K": "#0277BD", "L": "#00695C",
    "M": "#558B2F", "N": "#1565C0", "O": "#4527A0", "P": "#283593",
    "Q": "#01579B", "R": "#BF360C", "S": "#827717", "T": "#37474F",
    "U": "#795548", "V": "#616161",
}


# ── DB queries ─────────────────────────────────────────────────────────────────

def get_section_counts(conn, repo_id):
    return conn.execute("""
        SELECT cv.section, cv.section_name, COUNT(*) AS cnt
        FROM projects p JOIN classifications_vote cv ON p.id = cv.project_id
        WHERE p.repository_id = ?
        GROUP BY cv.section ORDER BY cnt DESC
    """, (repo_id,)).fetchall()


def get_type_stats(conn, repo_id):
    return conn.execute("""
        SELECT COALESCE(type, 'UNKNOWN'), COUNT(*) AS cnt
        FROM projects WHERE repository_id = ?
        GROUP BY type ORDER BY cnt DESC
    """, (repo_id,)).fetchall()


def get_top20(conn, repo_id):
    return conn.execute("""
        SELECT cv.section, cv.section_name,
               COALESCE(cv.division, '') AS division,
               COALESCE(cv.division_name, '') AS division_name,
               COUNT(*) AS cnt,
               ROUND(COUNT(*) * 100.0 / (
                   SELECT COUNT(*) FROM projects p2
                   JOIN classifications_vote cv2 ON p2.id = cv2.project_id
                   WHERE p2.repository_id = ?
               ), 1) AS pct
        FROM projects p JOIN classifications_vote cv ON p.id = cv.project_id
        WHERE p.repository_id = ?
        GROUP BY cv.section, cv.division
        ORDER BY cnt DESC LIMIT 20
    """, (repo_id, repo_id)).fetchall()


def get_confidence_stats(conn, repo_id):
    return conn.execute("""
        SELECT cv.confidence, COUNT(*) AS cnt
        FROM projects p JOIN classifications_vote cv ON p.id = cv.project_id
        WHERE p.repository_id = ?
        GROUP BY cv.confidence
        ORDER BY CASE cv.confidence
            WHEN 'very_high' THEN 1 WHEN 'high' THEN 2
            WHEN 'medium' THEN 3 ELSE 4 END
    """, (repo_id,)).fetchall()


def get_top_divisions(conn, repo_id, n=15):
    return conn.execute("""
        SELECT cv.section, cv.division,
               COALESCE(cv.division_name, cv.section_name, '') AS label,
               COUNT(*) AS cnt
        FROM projects p JOIN classifications_vote cv ON p.id = cv.project_id
        WHERE p.repository_id = ?
          AND cv.division IS NOT NULL AND cv.division != ''
        GROUP BY cv.section, cv.division
        ORDER BY cnt DESC LIMIT ?
    """, (repo_id, n)).fetchall()


def get_file_stats(conn, repo_id):
    """Return (total_files, conf_rows, top_section_rows) for file-level classifications."""
    total = conn.execute("""
        SELECT COUNT(*) FROM classifications_files cf
        JOIN projects p ON p.id = cf.project_id
        WHERE p.repository_id = ?
    """, (repo_id,)).fetchone()[0]
    if total == 0:
        return 0, [], []
    conf_rows = conn.execute("""
        SELECT cf.confidence, COUNT(*) as n,
               ROUND(COUNT(*)*100.0/?, 1) as pct
        FROM classifications_files cf JOIN projects p ON p.id = cf.project_id
        WHERE p.repository_id = ?
        GROUP BY cf.confidence
        ORDER BY CASE cf.confidence WHEN 'very_high' THEN 1 WHEN 'high' THEN 2
                 WHEN 'medium' THEN 3 ELSE 4 END
    """, (total, repo_id)).fetchall()
    top_rows = conn.execute("""
        SELECT cf.section, cf.section_name, COUNT(*) as n,
               ROUND(COUNT(*)*100.0/?, 1) as pct
        FROM classifications_files cf JOIN projects p ON p.id = cf.project_id
        WHERE p.repository_id = ?
        GROUP BY cf.section ORDER BY n DESC LIMIT 10
    """, (total, repo_id)).fetchall()
    return total, conf_rows, top_rows


# ── Chart generators → in-memory PNG bytes ─────────────────────────────────────

def chart_section_histogram(section_counts, type_stats, conf_stats,
                             repo_id, repo_name) -> bytes:
    total    = sum(c for _, _, c in section_counts)
    sec_dict = {s: c for s, _, c in section_counts}
    all_secs = sorted(ISIC_SECTIONS.keys())
    counts   = [sec_dict.get(s, 0) for s in all_secs]
    top3     = {s for s, _ in sorted(sec_dict.items(), key=lambda x: -x[1])[:3]}
    colors   = [BLUE_DARK if s in top3 else BLUE_LIGHT for s in all_secs]

    fig, ax = plt.subplots(figsize=(13, 5.5))
    fig.patch.set_facecolor("white")

    bars = ax.bar(range(len(all_secs)), counts, color=colors,
                  edgecolor="white", linewidth=0.4, width=0.75)

    max_count = max(counts) if counts else 1
    for bar, count, sec in zip(bars, counts, all_secs):
        if count > 0:
            ax.text(bar.get_x() + bar.get_width() / 2,
                    bar.get_height() + max_count * 0.012,
                    f"{count:,}", ha="center", va="bottom",
                    fontsize=6.5,
                    fontweight="bold" if sec in top3 else "normal",
                    color="#111" if sec in top3 else "#444")

    ax.set_xlim(-0.6, len(all_secs) - 0.4)
    ax.set_ylim(0, max_count * 1.18)
    ax.set_xticks(range(len(all_secs)))
    ax.set_xticklabels([f"{s}\n{SHORT_LABELS.get(s,'')}" for s in all_secs], fontsize=7)
    ax.set_xlabel("ISIC Rev.5 Section", fontsize=11, labelpad=8)
    ax.set_ylabel("Number of Projects", fontsize=11)
    ax.set_title(f"Primary ISIC Sections — Repository {repo_id}: {repo_name}  "
                 f"({total:,} projects)", fontsize=13, pad=10, fontweight="bold")
    ax.yaxis.grid(True, linestyle="--", alpha=0.6, color=GRID_COLOR)
    ax.set_axisbelow(True)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)

    conf_text = "   ".join(f"{c}: {n:,}" for c, n in conf_stats)
    ax.annotate(f"Confidence — {conf_text}", xy=(0.01, 0.97),
                xycoords="axes fraction", fontsize=8, va="top", color="#555")

    legend_elements = [
        Patch(facecolor=BLUE_DARK,  label="Top-3 sections"),
        Patch(facecolor=BLUE_LIGHT, label="Other sections"),
    ]
    ax.legend(handles=legend_elements, loc="upper right", fontsize=8, framealpha=0.7)

    fig.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf.read()


def chart_division_bars(top_divs, section_counts, repo_id, repo_name) -> bytes:
    if not top_divs:
        return b""
    total  = sum(c for _, _, c in section_counts)
    labels = []
    values = []
    colors = []
    for sec, div, label, cnt in reversed(top_divs):
        short = label[:48] if len(label) > 48 else label
        labels.append(f"{sec}/{div}  {short}")
        values.append(cnt)
        colors.append(_SEC_COLORS.get(sec, BLUE_DARK))

    fig, ax = plt.subplots(figsize=(13, 6.5))
    fig.patch.set_facecolor("white")

    y_pos = range(len(labels))
    bars  = ax.barh(list(y_pos), values, color=colors,
                    edgecolor="white", linewidth=0.4, height=0.7)
    max_val = max(values) if values else 1
    for bar, val in zip(bars, values):
        ax.text(bar.get_width() + max_val * 0.008,
                bar.get_y() + bar.get_height() / 2,
                f"{val:,}", va="center", ha="left", fontsize=7.5, color="#222")

    ax.set_yticks(list(y_pos))
    ax.set_yticklabels(labels, fontsize=8)
    ax.set_xlim(0, max_val * 1.18)
    ax.set_xlabel("Number of Projects", fontsize=11)
    ax.set_title(f"Top-{len(top_divs)} ISIC Divisions — Repository {repo_id}: {repo_name}",
                 fontsize=13, pad=12, fontweight="bold")
    ax.xaxis.grid(True, linestyle="--", alpha=0.5, color=GRID_COLOR)
    ax.set_axisbelow(True)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)

    seen_secs = list(dict.fromkeys(sec for sec, *_ in top_divs))
    legend_handles = [
        Patch(facecolor=_SEC_COLORS.get(s, BLUE_DARK),
              label=f"{s} — {ISIC_SECTIONS.get(s, s)[:35]}")
        for s in seen_secs
    ]
    ax.legend(handles=legend_handles, loc="lower right", fontsize=7,
              framealpha=0.85, ncol=2)

    fig.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf.read()


# ── Word document helpers ───────────────────────────────────────────────────────

def _rgb(r, g, b): return RGBColor(r, g, b)

BLUE_DARK_RGB  = (0x15, 0x65, 0xC0)
ALT_ROW_RGB    = (0xE3, 0xF2, 0xFD)


def set_cell_bg(cell, rgb_tuple):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  f"{rgb_tuple[0]:02X}{rgb_tuple[1]:02X}{rgb_tuple[2]:02X}")
    tcPr.append(shd)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(15 if level == 1 else 12)
    run.font.color.rgb = _rgb(*BLUE_DARK_RGB)
    if level == 1:
        pPr  = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot  = OxmlElement("w:bottom")
        bot.set(qn("w:val"),   "single")
        bot.set(qn("w:sz"),    "6")
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), "1565C0")
        pBdr.append(bot)
        pPr.append(pBdr)


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(5)
    p.runs[0].font.size = Pt(10.5)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.add_run(text).font.size = Pt(10.5)


def add_chart_image(doc, png_bytes: bytes, width_inches: float = 6.3):
    if not png_bytes:
        return
    buf = io.BytesIO(png_bytes)
    doc.add_picture(buf, width=Inches(width_inches))
    # Center the picture paragraph
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_top20_table(doc, top20):
    headers = ["Rank", "Sec.", "Division", "Class Name", "Count", "%"]
    tbl = doc.add_table(rows=1 + len(top20), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = "Table Grid"

    col_widths_cm = [1.1, 1.1, 1.8, 9.5, 1.8, 1.4]

    # Header row
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.width = Cm(col_widths_cm[i])
        set_cell_bg(cell, BLUE_DARK_RGB)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p   = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)
        p.paragraph_format.space_after  = Pt(1)
        p.paragraph_format.space_before = Pt(1)

    # Data rows
    for r_idx, (sec, sec_name, div, div_name, cnt, pct) in enumerate(top20):
        row = tbl.rows[r_idx + 1]
        if r_idx % 2 == 0:
            for cell in row.cells:
                set_cell_bg(cell, ALT_ROW_RGB)

        class_label = (f"{div} — {div_name}" if div and div_name
                       else div_name or div or sec_name or "")
        values = [str(r_idx + 1), sec, div or "—", class_label[:80],
                  f"{cnt:,}", f"{pct:.1f}%"]

        for c_idx, val in enumerate(values):
            cell = row.cells[c_idx]
            cell.width = Cm(col_widths_cm[c_idx])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p   = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(9)
            p.paragraph_format.space_after  = Pt(1)
            p.paragraph_format.space_before = Pt(1)

    doc.add_paragraph()  # spacer


def add_section_summary_table(doc, section_counts, total):
    add_heading(doc, "Section Distribution", level=2)
    headers = ["Section", "Name", "Projects", "%"]
    tbl = doc.add_table(rows=1 + len(section_counts), cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = "Table Grid"

    col_widths_cm = [1.5, 9.0, 2.5, 1.8]

    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.width = Cm(col_widths_cm[i])
        set_cell_bg(cell, BLUE_DARK_RGB)
        p   = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)
        p.paragraph_format.space_after  = Pt(1)
        p.paragraph_format.space_before = Pt(1)

    for r_idx, (sec, sec_name, cnt) in enumerate(section_counts):
        row = tbl.rows[r_idx + 1]
        if r_idx % 2 == 0:
            for cell in row.cells:
                set_cell_bg(cell, ALT_ROW_RGB)
        values = [sec, sec_name or ISIC_SECTIONS.get(sec, ""),
                  f"{cnt:,}", f"{cnt*100/total:.1f}%"]
        for c_idx, val in enumerate(values):
            cell = row.cells[c_idx]
            cell.width = Cm(col_widths_cm[c_idx])
            p = cell.paragraphs[0]
            p.add_run(val).font.size = Pt(9)
            p.paragraph_format.space_after  = Pt(1)
            p.paragraph_format.space_before = Pt(1)

    doc.add_paragraph()


# ── Build comments text ─────────────────────────────────────────────────────────

def build_comments(repo_id, repo_name, section_counts, type_stats, conf_stats):
    total      = sum(c for _, _, c in section_counts)
    sec_dict   = {s: c for s, _, c in section_counts}
    sorted_secs = sorted(sec_dict.items(), key=lambda x: -x[1])
    dom_sec, dom_cnt = sorted_secs[0] if sorted_secs else ("?", 0)
    dom_name  = ISIC_SECTIONS.get(dom_sec, "")
    dom_pct   = dom_cnt * 100.0 / total if total else 0.0
    n_secs    = len(sec_dict)
    very_high = sum(n for c, n in conf_stats if c == "very_high")
    high      = sum(n for c, n in conf_stats if c == "high")
    confident = very_high + high
    conf_pct  = confident * 100.0 / total if total else 0.0

    type_str = ";  ".join(f"{t} ({c:,})" for t, c in type_stats)
    secs_list = ", ".join(sorted(sec_dict.keys()))

    bullets = [
        f"Repository {repo_id} ({repo_name}) contains {total:,} classified projects "
        f"across {n_secs} ISIC Rev.5 sections ({secs_list}).",
        f"Project types present: {type_str}.",
        f"The dominant ISIC section is {dom_sec} — {dom_name}, accounting for "
        f"{dom_cnt:,} projects ({dom_pct:.1f}% of the repository).",
        f"Classification confidence: {confident:,} projects ({conf_pct:.1f}%) "
        f"classified at high or very-high confidence by the voting ensemble.",
    ]

    if dom_sec == "N":
        bullets.append(
            "Section N (Professional, scientific and technical activities) dominates. "
            "This is expected for general-purpose academic repositories — many datasets "
            "describe scientific research processes. Division 72 (Scientific research and "
            "development) is typically the most frequent sub-class."
        )
    elif dom_sec == "A":
        bullets.append(
            "Section A (Agriculture, forestry and fishing) dominates, indicating a strong "
            "representation of environmental, ecological, and agricultural datasets — consistent "
            "with repositories serving the natural-sciences community."
        )
    elif dom_sec == "R":
        bullets.append(
            "Section R (Human health and social work activities) is the dominant class, "
            "reflecting the repository's focus on medical, clinical, and social datasets."
        )
    elif dom_sec == "P":
        bullets.append(
            "Section P (Public administration and defence) dominates, reflecting a large "
            "share of political-science, governance, and survey datasets typical of social "
            "science repositories."
        )
    elif dom_sec == "Q":
        bullets.append(
            "Section Q (Education) is dominant, consistent with repositories that focus on "
            "educational research datasets."
        )

    if n_secs >= 18:
        bullets.append(
            f"The distribution spans {n_secs} of 22 ISIC sections, indicating high thematic "
            "diversity. No single field monopolises the repository."
        )
    elif n_secs <= 8:
        bullets.append(
            f"With only {n_secs} sections represented, the repository has a narrow thematic "
            "focus. Future curation may benefit from broadening the scope."
        )

    return bullets


# ── Main report builder ─────────────────────────────────────────────────────────

def generate_repo_docx(conn, repo_id: int) -> str:
    repo_name      = REPO_NAMES.get(repo_id, f"Repository {repo_id}")
    section_counts = get_section_counts(conn, repo_id)
    type_stats     = get_type_stats(conn, repo_id)
    top20          = get_top20(conn, repo_id)
    conf_stats     = get_confidence_stats(conn, repo_id)
    top_divs       = get_top_divisions(conn, repo_id, n=15)
    total          = sum(c for _, _, c in section_counts)
    file_total, file_conf, file_top_secs = get_file_stats(conn, repo_id)

    doc = Document()

    # Page margins
    for sec in doc.sections:
        sec.top_margin    = Cm(2.2)
        sec.bottom_margin = Cm(2.2)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.5)

    # ── Title block ──────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after  = Pt(4)
    r = title.add_run("SQ26 — ISIC Rev.5 Classification Report")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = _rgb(*BLUE_DARK_RGB)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(2)
    r2 = sub.add_run(f"Repository {repo_id}: {repo_name}  ·  {total:,} projects")
    r2.font.size = Pt(13)
    r2.font.color.rgb = _rgb(0x55, 0x55, 0x55)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(16)
    r3 = meta.add_run(
        "Student ID: 23692652  ·  FAU Erlangen-Nürnberg  ·  SQ26 Project  ·  July 2026"
    )
    r3.font.size = Pt(9)
    r3.font.color.rgb = _rgb(0xAA, 0xAA, 0xAA)

    # ── Section 1: Histogram ─────────────────────────────────────────
    add_heading(doc, "1. Primary ISIC Section Distribution")
    hist_png = chart_section_histogram(
        section_counts, type_stats, conf_stats, repo_id, repo_name
    )
    add_chart_image(doc, hist_png, width_inches=6.0)
    doc.add_paragraph()

    # ── Section 2: Full section table ───────────────────────────────
    add_section_summary_table(doc, section_counts, total)

    # ── Section 3: Top-20 table ──────────────────────────────────────
    add_heading(doc, "2. Top-20 ISIC Classes (Section + Division)")
    add_top20_table(doc, top20)

    # ── Section 4: Comments ──────────────────────────────────────────
    add_heading(doc, "3. Comments on Findings")
    comments = build_comments(repo_id, repo_name, section_counts, type_stats, conf_stats)
    for bullet_text in comments:
        add_bullet(doc, bullet_text)
    doc.add_paragraph()

    # ── Section 5: Division bar chart ────────────────────────────────
    add_heading(doc, "4. Top-15 ISIC Divisions (Secondary Classes)")
    div_png = chart_division_bars(top_divs, section_counts, repo_id, repo_name)
    add_chart_image(doc, div_png, width_inches=6.0)

    # ── Section 6: File-level classification (if available) ──────────
    if file_total > 0:
        add_heading(doc, "5. File-Level ISIC Classification")
        add_body(doc, (
            f"Individual files within QDA_PROJECT and QD_PROJECT were classified "
            f"using the fine-tuned bert-base-uncased model (bert_base_isic_v1). "
            f"Text was constructed from the parent project title, description, "
            f"keywords, file name, and the stored file text snippet. "
            f"Repository {repo_id} contributed {file_total:,} classified files."
        ))

        # Confidence table
        add_body(doc, "Confidence distribution:")
        conf_tbl = doc.add_table(rows=1 + len(file_conf), cols=3)
        conf_tbl.style = "Table Grid"
        hdr = conf_tbl.rows[0].cells
        for i, h in enumerate(["Confidence", "Files", "%"]):
            hdr[i].text = h
            hdr[i].paragraphs[0].runs[0].bold = True
            set_cell_bg(hdr[i], BLUE_DARK_RGB)
            hdr[i].paragraphs[0].runs[0].font.color.rgb = _rgb(0xFF, 0xFF, 0xFF)
        for ri, (conf, n, pct) in enumerate(file_conf, start=1):
            row = conf_tbl.rows[ri].cells
            row[0].text = conf
            row[1].text = f"{n:,}"
            row[2].text = f"{pct}%"
        doc.add_paragraph()

        # Top sections table
        add_body(doc, "Top ISIC sections across classified files:")
        sec_tbl = doc.add_table(rows=1 + len(file_top_secs), cols=4)
        sec_tbl.style = "Table Grid"
        hdr2 = sec_tbl.rows[0].cells
        for i, h in enumerate(["Sec", "Name", "Files", "%"]):
            hdr2[i].text = h
            hdr2[i].paragraphs[0].runs[0].bold = True
            set_cell_bg(hdr2[i], BLUE_DARK_RGB)
            hdr2[i].paragraphs[0].runs[0].font.color.rgb = _rgb(0xFF, 0xFF, 0xFF)
        alt = (0xDE, 0xEA, 0xF1)
        for ri, (sec, sec_name, n, pct) in enumerate(file_top_secs, start=1):
            row = sec_tbl.rows[ri].cells
            row[0].text = sec
            row[1].text = (sec_name or "")[:52]
            row[2].text = f"{n:,}"
            row[3].text = f"{pct}%"
            if ri % 2 == 0:
                for cell in row:
                    set_cell_bg(cell, alt)
        doc.add_paragraph()

    out_path = f"23692652-sq26-report-repo{repo_id}.docx"
    doc.save(out_path)
    print(f"[repo {repo_id}] Saved: {out_path}  ({total:,} projects)")
    return out_path


def main() -> None:
    ap = argparse.ArgumentParser(
        description="Generate Word (.docx) classification reports per repository."
    )
    ap.add_argument("--db",   default=DB)
    ap.add_argument("--repo", nargs="+", type=int, default=None)
    args = ap.parse_args()

    conn = sqlite3.connect(args.db)
    repo_ids = args.repo or [
        r[0] for r in conn.execute(
            "SELECT DISTINCT repository_id FROM projects ORDER BY repository_id"
        ).fetchall()
    ]

    print(f"Generating Word reports for {len(repo_ids)} repositories…\n")
    for rid in repo_ids:
        generate_repo_docx(conn, rid)

    conn.close()
    print("\nAll done.")


if __name__ == "__main__":
    main()
