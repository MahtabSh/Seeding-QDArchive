#!/usr/bin/env python3
"""Convert 23692652-sq26-classification-methodology.md to a styled Word document."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
from pathlib import Path

SRC  = "23692652-sq26-classification-methodology.md"
OUT  = "23692652-sq26-classification-methodology.docx"

BLUE_DARK  = (0x1F, 0x4E, 0x79)   # header fill
BLUE_MID   = (0x27, 0x63, 0xA0)   # h1 text
BLUE_LIGHT = (0x22, 0x96, 0xF3)   # accent
ALT_ROW    = (0xDE, 0xEA, 0xF1)   # alternating table row

def _rgb(t): return RGBColor(t[0], t[1], t[2])


# ── helpers ────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, rgb) -> None:
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def set_cell_border(cell, **kwargs) -> None:
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{side}")
        tag.set(qn("w:val"),   kwargs.get("val",   "single"))
        tag.set(qn("w:sz"),    kwargs.get("sz",    "4"))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), kwargs.get("color", "AAAAAA"))
        tcBorders.append(tag)
    tcPr.append(tcBorders)


def add_paragraph_style(doc: Document) -> None:
    """Ensure Normal style has adequate spacing."""
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    pf = style.paragraph_format
    pf.space_after  = Pt(4)
    pf.space_before = Pt(0)


def h1(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold       = True
    run.font.size  = Pt(16)
    run.font.color.rgb = _rgb(BLUE_MID)
    # bottom border
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "2763A0")
    pBdr.append(bot)
    pPr.append(pBdr)


def h2(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(13)
    run.font.color.rgb = _rgb(BLUE_DARK)


def h3(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(9)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(11.5)
    run.font.color.rgb = RGBColor(0x37, 0x47, 0x4F)


def body(doc: Document, text: str) -> None:
    if not text.strip():
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    _inline(p, text)


def bullet(doc: Document, text: str, level: int = 0) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent  = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after  = Pt(2)
    _inline(p, text)


def code_block(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x37, 0x47, 0x4F)


def _inline(p, text: str) -> None:
    """Parse **bold**, `code`, and plain text within a line."""
    pattern = re.compile(r'(\*\*[^*]+\*\*|`[^`]+`)')
    parts   = pattern.split(text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)
        else:
            p.add_run(part)


def md_table(doc: Document, header_row: list[str], rows: list[list[str]]) -> None:
    col_count = len(header_row)
    tbl = doc.add_table(rows=1 + len(rows), cols=col_count)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = "Table Grid"

    # Header
    hdr = tbl.rows[0]
    for i, text in enumerate(header_row):
        cell = hdr.cells[i]
        set_cell_bg(cell, BLUE_DARK)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p   = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold           = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size      = Pt(9.5)
        p.paragraph_format.space_after  = Pt(2)
        p.paragraph_format.space_before = Pt(2)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = tbl.rows[r_idx + 1]
        bg  = _rgb(ALT_ROW) if r_idx % 2 == 0 else None
        for c_idx, text in enumerate(row_data):
            cell = row.cells[c_idx]
            if bg:
                set_cell_bg(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            _inline(p, text)
            p.runs[0].font.size = Pt(9.5) if p.runs else None
            p.paragraph_format.space_after  = Pt(2)
            p.paragraph_format.space_before = Pt(2)

    doc.add_paragraph()  # spacer after table


# ── main converter ─────────────────────────────────────────────────────────────

def convert(src: str, out: str) -> None:
    lines = Path(src).read_text(encoding="utf-8").splitlines()
    doc   = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    add_paragraph_style(doc)

    # Title page block (first 6 lines are front matter)
    title_done = False
    i = 0
    n = len(lines)

    # Collect markdown table rows across lines
    table_lines: list[str] = []
    in_code = False
    code_lines: list[str] = []

    while i < n:
        line = lines[i]

        # ── Fenced code block ──
        if line.strip().startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                in_code = False
                code_block(doc, "\n".join(code_lines))
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # ── Horizontal rule ──
        if re.match(r"^---+$", line.strip()):
            p = doc.add_paragraph()
            pPr  = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bot  = OxmlElement("w:bottom")
            bot.set(qn("w:val"),   "single")
            bot.set(qn("w:sz"),    "4")
            bot.set(qn("w:space"), "1")
            bot.set(qn("w:color"), "CCCCCC")
            pBdr.append(bot)
            pPr.append(pBdr)
            i += 1
            continue

        # ── Headings ──
        m = re.match(r"^(#{1,3})\s+(.*)", line)
        if m:
            level = len(m.group(1))
            text  = m.group(2).strip()
            if level == 1:
                h1(doc, text)
            elif level == 2:
                h2(doc, text)
            else:
                h3(doc, text)
            i += 1
            continue

        # ── Markdown table ──
        if line.strip().startswith("|"):
            # Gather all consecutive table lines
            tbl_raw = []
            while i < n and lines[i].strip().startswith("|"):
                tbl_raw.append(lines[i])
                i += 1
            # Parse: first line = header, second = separator, rest = data
            def parse_row(r):
                return [c.strip() for c in r.strip().strip("|").split("|")]
            header = parse_row(tbl_raw[0])
            data   = [parse_row(r) for r in tbl_raw[2:]]  # skip separator
            md_table(doc, header, data)
            continue

        # ── Bullet list ──
        bm = re.match(r"^(\s*)[-*]\s+(.*)", line)
        if bm:
            indent = len(bm.group(1)) // 2
            bullet(doc, bm.group(2), level=indent)
            i += 1
            continue

        # ── Bold metadata lines (** key:** value) at top ──
        if line.startswith("**") and not title_done:
            p = doc.add_paragraph()
            _inline(p, line)
            i += 1
            continue

        # ── Empty line ──
        if not line.strip():
            i += 1
            continue

        # ── Normal paragraph ──
        body(doc, line)
        i += 1

    doc.save(out)
    print(f"Saved → {out}")


if __name__ == "__main__":
    convert(SRC, OUT)
